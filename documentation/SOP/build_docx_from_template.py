"""
Build DOCX from Markdown under the SOP folder. Published SOP Markdown lives on SharePoint
(site HTOCDataAnalyticsASA) at Documents/HTOC Data Analytics/SOPs/ — keep this script beside
the synced .md sources (or set paths below).
"""
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SOP_DIR = Path(__file__).resolve().parent
# Override with absolute path via SOP_DOCX_TEMPLATE when the template is not beside this script.
TEMPLATE = Path(
    os.environ.get("SOP_DOCX_TEMPLATE", str(SOP_DIR / "SOP_Format_Example.docx"))
)
OUT_DIR = SOP_DIR / "DOCX_v11"
OUT_DIR.mkdir(exist_ok=True)


def add_inline_markdown_runs(paragraph, text: str):
    """
    Render simple inline markdown styling:
    - **bold**
    - `inline code`
    Everything else is emitted as normal text.
    """
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    path_pattern = re.compile(r"([A-Za-z]:\\[^\s,;)\]]+|\\\\[^\s,;)\]]+|/[A-Za-z0-9._-]+(?:/[A-Za-z0-9._\-\s]+)+)")

    def add_plain_with_path_emphasis(segment: str):
        last = 0
        for m in path_pattern.finditer(segment):
            if m.start() > last:
                paragraph.add_run(segment[last:m.start()])
            run = paragraph.add_run(m.group(0))
            run.bold = True
            run.italic = True
            run.underline = True
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            last = m.end()
        if last < len(segment):
            paragraph.add_run(segment[last:])

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            code_text = part[1:-1]
            run = paragraph.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            if "\\" in code_text or "/" in code_text:
                run.bold = True
                run.underline = True
                run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        else:
            add_plain_with_path_emphasis(part)


def remove_after_cover(doc: Document):
    """
    Keep template cover content/elements and remove body paragraphs from first
    'Purpose' heading onward.
    """
    start_para = None
    for p in doc.paragraphs:
        if p.text.strip() == "Purpose":
            start_para = p
            break

    if start_para is None:
        return

    body = doc._element.body
    removing = False
    for child in list(body):
        if child.tag == qn('w:p') and child is start_para._p:
            removing = True
        if removing and child.tag != qn('w:sectPr'):
            body.remove(child)

    # Remove leftover blank page-break paragraphs retained from template.
    for p in list(doc.paragraphs):
        xml = p._p.xml
        has_page_break = 'w:type="page"' in xml
        if has_page_break and not p.text.strip():
            p._p.getparent().remove(p._p)


def update_cover_text(doc: Document, subtitle: str, version_text: str, last_reviewed: str):
    # Replace template cover subtitle and version line while preserving cover graphics.
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "Insight Report Distribution Procedures":
            p.text = subtitle
        elif "Version 1.0" in t:
            p.text = version_text

    # Replace template header title text with project/SOP name
    header_text = f"TLP:AMBER\n{subtitle}"
    if last_reviewed:
        header_text += f"\n{last_reviewed}"

    sec = doc.sections[0]
    for tbl in sec.header.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if "Insight Report Distribution Procedures" in cell.text:
                    cell.text = header_text


def remove_existing_toc_paragraphs(doc: Document):
    # Remove any TOC field placeholders inherited from the template.
    for p in list(doc.paragraphs):
        xml = p._p.xml
        if 'TOC \\o' in xml or "Right-click to update field." in p.text:
            p._p.getparent().remove(p._p)


def flush_table(doc, table_lines):
    if not table_lines:
        return

    rows = []
    for ln in table_lines:
        parts = [p.strip() for p in ln.strip().strip("|").split("|")]
        rows.append(parts)

    if len(rows) < 2:
        for ln in table_lines:
            p = doc.add_paragraph("")
            add_inline_markdown_runs(p, ln)
        return

    header = rows[0]
    data_rows = rows[2:] if len(rows) >= 2 else []
    cols = len(header)
    t = doc.add_table(rows=1, cols=cols)
    t.style = "Table Grid"

    for i, h in enumerate(header):
        t.rows[0].cells[i].text = h

    for dr in data_rows:
        r = t.add_row().cells
        for i in range(cols):
            r[i].text = dr[i] if i < len(dr) else ""

    # Improve pagination behavior: avoid broken-looking tables.
    # 1) Keep the paragraph right before a table with the table.
    if len(doc.paragraphs) >= 1:
        prev = doc.paragraphs[-1]
        try:
            prev.paragraph_format.keep_with_next = True
        except Exception:
            pass

    # 2) Repeat header row on each page and avoid splitting table rows.
    for idx, row in enumerate(t.rows):
        tr = row._tr
        trPr = tr.get_or_add_trPr()

        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)

        if idx == 0:
            tbl_header = OxmlElement("w:tblHeader")
            trPr.append(tbl_header)


def add_markdown_content(doc: Document, lines):
    style_names = {s.name for s in doc.styles}
    list_number_style = "List Number" if "List Number" in style_names else "Normal"
    list_bullet_style = "List Bullet" if "List Bullet" in style_names else "Normal"
    code_style = "No Spacing" if "No Spacing" in style_names else "Normal"

    in_code = False
    table_buf = []

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            flush_table(doc, table_buf)
            table_buf = []
            in_code = not in_code
            continue

        if in_code:
            p = doc.add_paragraph(line)
            p.style = code_style
            if p.runs:
                p.runs[0].font.name = "Consolas"
                p.runs[0].font.size = Pt(9)
            continue

        if line.startswith("|"):
            table_buf.append(line)
            continue
        else:
            flush_table(doc, table_buf)
            table_buf = []

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.strip() == "---":
            doc.add_paragraph("-")
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 6)
            text = m.group(2).strip()
            p = doc.add_heading("", level=level)
            add_inline_markdown_runs(p, text)
            continue

        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph("", style=list_number_style)
            add_inline_markdown_runs(p, text)
            continue

        if line.lstrip().startswith("- "):
            text = line.lstrip()[2:]
            p = doc.add_paragraph("", style=list_bullet_style)
            add_inline_markdown_runs(p, text)
            continue

        p = doc.add_paragraph("")
        add_inline_markdown_runs(p, line)

    flush_table(doc, table_buf)


def add_word_toc_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instr)

    run = paragraph.add_run()
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    paragraph.add_run("Right-click and update field.")

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_doc_specific_toc(doc: Document):
    # Cover -> TOC -> body
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    toc_para = doc.add_paragraph("")
    add_word_toc_field(toc_para)
    doc.add_page_break()


def extract_body_lines(md_lines):
    # Skip the initial metadata table block and start at first section after divider.
    idx = 0
    while idx < len(md_lines):
        if md_lines[idx].strip() == "---":
            return md_lines[idx + 1 :]
        idx += 1
    return md_lines


def parse_meta(md_lines):
    h2 = ""
    version = ""
    last_reviewed = ""

    for ln in md_lines:
        if ln.startswith("## ") and not h2:
            h2 = ln[3:].strip()
        if ln.startswith("| **Version** "):
            version = ln.split("|", 3)[2].strip() if len(ln.split("|")) > 2 else ""
        if ln.startswith("| **Last Reviewed** "):
            parts = ln.split("|")
            if len(parts) >= 4:
                last_reviewed = parts[2].strip()

    if not version:
        version = "1.0"
    if not last_reviewed:
        last_reviewed = ""

    return (
        h2 or "Standard Operating Procedure",
        f"Version {version}\n{last_reviewed}".strip(),
        last_reviewed,
    )


def build_one(md_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    body_lines = extract_body_lines(lines)

    subtitle, version_text, last_reviewed = parse_meta(lines)

    doc = Document(str(TEMPLATE))
    remove_after_cover(doc)
    remove_existing_toc_paragraphs(doc)
    update_cover_text(doc, subtitle, version_text, last_reviewed)

    # Insert a single doc-specific TOC for this SOP.
    add_doc_specific_toc(doc)
    add_markdown_content(doc, body_lines)

    out = OUT_DIR / (md_path.stem + ".docx")
    doc.save(str(out))
    print(f"Created: {out}")


def main():
    for md in sorted(SOP_DIR.glob("SOP_*.md")):
        build_one(md)


if __name__ == "__main__":
    main()
